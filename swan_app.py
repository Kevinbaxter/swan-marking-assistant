import streamlit as st
from docx import Document
import re
from collections import Counter
import random

SHORT_PARAGRAPH_LIMIT = 260

# -----------------------------
# TEXT EXTRACTION
# -----------------------------
def extract_text(doc):
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return paragraphs

def count_headings(doc):
    return sum(1 for p in doc.paragraphs if p.style.name.startswith("Heading"))

def has_bullets(doc):
    return any("List" in p.style.name for p in doc.paragraphs)

# -----------------------------
# BASIC FEATURES
# -----------------------------
def find_short_paragraphs(paragraphs):
    return [p for p in paragraphs if len(p) < SHORT_PARAGRAPH_LIMIT]

def find_spelling_issues(paragraphs):
    words = re.findall(r"\b[a-zA-Z]{3,}\b", " ".join(paragraphs).lower())
    counts = Counter(words)
    # crude heuristic: repeated 3+ times might be a misspelling
    return [w for w, c in counts.items() if c >= 3 and w not in COMMON_WORDS]

COMMON_WORDS = set("""
the and a an in on at of to for with from by is are was were be been being
this that these those it they them he she we you i as if or but so because
""".split())

# -----------------------------
# ADVANCED ANALYSIS
# -----------------------------
def sentence_variety(paragraphs):
    text = " ".join(paragraphs)
    sentences = re.split(r"[.!?]", text)
    lengths = [len(s.split()) for s in sentences if len(s.split()) > 0]
    if not lengths:
        return None, None
    avg = sum(lengths) / len(lengths)
    return avg, lengths

def vocabulary_richness(paragraphs):
    words = re.findall(r"\b[a-zA-Z]{3,}\b", " ".join(paragraphs).lower())
    if not words:
        return 0
    unique = set(words)
    return len(unique) / len(words)

LINKERS = [
    "however", "therefore", "in addition", "furthermore", "moreover",
    "for example", "for instance", "consequently", "as a result", "on the other hand"
]

def has_linkers(paragraphs):
    text = " ".join(paragraphs).lower()
    return any(l in text for l in LINKERS)

def has_argument_structure(paragraphs):
    text = " ".join(paragraphs).lower()
    markers = ["because", "this shows", "this suggests", "therefore", "as a result"]
    return any(m in text for m in markers)

def has_clear_conclusion(paragraphs):
    if not paragraphs:
        return False
    last = paragraphs[-1].lower()
    return any(phrase in last for phrase in ["in conclusion", "overall", "to sum up", "in summary"])

# -----------------------------
# FEEDBACK TEMPLATES
# -----------------------------
STRENGTH_TEMPLATES = [
    "Your writing shows clear control of {feature}.",
    "You demonstrate strong use of {feature}.",
    "A real strength in your work is your {feature}.",
    "You consistently use {feature}, which supports your ideas well.",
]

WEAKNESS_TEMPLATES = [
    "Your writing would be stronger with more focus on {feature}.",
    "At times, your work lacks {feature}.",
    "There is room to improve your use of {feature}.",
    "Currently, {feature} is not used consistently in your work.",
]

ACTION_TEMPLATES = [
    "Next time, try to {action}.",
    "As a next step, you could {action}.",
    "To improve, focus on {action}.",
    "A useful target would be to {action}.",
]

def strength_feedback(feature):
    return random.choice(STRENGTH_TEMPLATES).format(feature=feature)

def weakness_feedback(feature):
    return random.choice(WEAKNESS_TEMPLATES).format(feature=feature)

def action_feedback(action):
    return random.choice(ACTION_TEMPLATES).format(action=action)

# -----------------------------
# STREAMLIT APP
# -----------------------------
st.title("🦢 SWAN Marking Assistant – Word (Enhanced)")

uploaded = st.file_uploader("Upload a Word document (.docx)", type="docx")

if uploaded:
    doc = Document(uploaded)
    paragraphs = extract_text(doc)

    strengths = []
    weaknesses = []
    actions = []

    # STRUCTURE: headings, bullets, paragraphs
    if count_headings(doc) >= 1:
        strengths.append(strength_feedback("headings to organise your ideas"))
    else:
        weaknesses.append(weakness_feedback("clear headings to guide the reader"))
        actions.append(action_feedback("add headings to show where each new idea or section begins"))

    if has_bullets(doc):
        strengths.append(strength_feedback("bullet points to break up information"))
    else:
        weaknesses.append(weakness_feedback("bullet points to make key points stand out"))
        actions.append(action_feedback("use bullet points for lists or key ideas"))

    short_paras = find_short_paragraphs(paragraphs)
    if not short_paras:
        strengths.append(strength_feedback("well-developed paragraphs with enough detail"))
    else:
        weaknesses.append(weakness_feedback("paragraph development – one or more paragraphs are very short"))
        actions.append(action_feedback("expand short paragraphs by adding examples, explanations or evidence"))

    # SENTENCE VARIETY
    avg_len, lengths = sentence_variety(paragraphs)
    if avg_len:
        if avg_len < 10:
            weaknesses.append(weakness_feedback("sentence variety – many sentences are very short"))
            actions.append(action_feedback("combine some short sentences to create more complex ones"))
        elif avg_len > 25:
            weaknesses.append(weakness_feedback("sentence control – some sentences are very long"))
            actions.append(action_feedback("split long sentences into two shorter ones to improve clarity"))
        else:
            strengths.append(strength_feedback("a good mix of shorter and longer sentences"))

    # VOCABULARY
    vocab_score = vocabulary_richness(paragraphs)
    if vocab_score > 0.4:
        strengths.append(strength_feedback("varied and precise vocabulary"))
    elif vocab_score < 0.25 and vocab_score > 0:
        weaknesses.append(weakness_feedback("varied vocabulary"))
        actions.append(action_feedback("experiment with more ambitious word choices and avoid repeating the same words"))

    # LINKING & COHERENCE
    if has_linkers(paragraphs):
        strengths.append(strength_feedback("linking words to connect your ideas"))
    else:
        weaknesses.append(weakness_feedback("linking words to show how ideas connect"))
        actions.append(action_feedback("use phrases like 'however', 'in addition', or 'for example' to guide the reader"))

    # ARGUMENT / EXPLANATION
    if has_argument_structure(paragraphs):
        strengths.append(strength_feedback("explanations that show how your evidence supports your points"))
    else:
        weaknesses.append(weakness_feedback("clear explanation of how evidence supports your points"))
        actions.append(action_feedback("add phrases like 'this shows that…' or 'this suggests that…' after your evidence"))

    # CONCLUSION
    if has_clear_conclusion(paragraphs):
        strengths.append(strength_feedback("a clear concluding paragraph that rounds off your response"))
    else:
        weaknesses.append(weakness_feedback("a clear conclusion at the end of your work"))
        actions.append(action_feedback("add a short conclusion that sums up your main points and links back to the question"))

    # SPELLING
    spelling_issues = find_spelling_issues(paragraphs)
    if spelling_issues:
        weaknesses.append("The same word seems to be repeated several times and may be misspelt.")
        actions.append(action_feedback("use spellcheck or read your work aloud to spot repeated spelling errors"))

    # -----------------------------
    # OUTPUT
    # -----------------------------
    st.subheader("Strengths")
    if strengths:
        for s in strengths[:5]:
            st.write("•", s)
    else:
        st.write("• No clear strengths detected – check the document content.")

    st.subheader("Weaknesses")
    if weaknesses:
        for w in weaknesses[:5]:
            st.write("•", w)
    else:
        st.write("• No obvious weaknesses detected based on the current checks.")

    st.subheader("Next Steps")
    if actions:
        for a in actions[:5]:
            st.write("•", a)
    else:
        st.write("• No specific next steps generated – consider setting your own target based on the feedback above.")

